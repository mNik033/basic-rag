import io
import os
import uuid
import blake3
import docx
import pypdf
from app.core.exceptions import EmptyDocumentError, UnsupportedFileTypeError
from app.domain.schemas import DocumentMetadata, ParsedDocument

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}


class DocumentParser:
    """Extracts text content and metadata from various file formats with BLAKE3 hashing."""

    @staticmethod
    def compute_hash(content_bytes: bytes) -> str:
        """Compute hexadecimal BLAKE3 digest of raw file bytes."""
        hasher = blake3.blake3()
        hasher.update(content_bytes)
        return hasher.hexdigest()

    def parse(self, filename: str, content_bytes: bytes) -> ParsedDocument:
        """Parse raw file bytes into a ParsedDocument domain model."""
        ext = os.path.splitext(filename)[1].lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                extension=ext,
                supported=sorted(list(SUPPORTED_EXTENSIONS)),
            )

        if ext == ".pdf":
            text = self._parse_pdf(content_bytes)
        elif ext == ".docx":
            text = self._parse_docx(content_bytes)
        elif ext in {".txt", ".md"}:
            text = self._parse_text(content_bytes)
        else:
            raise UnsupportedFileTypeError(
                extension=ext,
                supported=sorted(list(SUPPORTED_EXTENSIONS)),
            )

        cleaned_text = text.strip()
        if not cleaned_text:
            raise EmptyDocumentError(filename=filename)

        content_hash = self.compute_hash(content_bytes)
        doc_id = str(uuid.uuid4())

        metadata = DocumentMetadata(
            filename=filename,
            content_hash=content_hash,
            file_type=ext.lstrip("."),
            file_size=len(content_bytes),
        )

        return ParsedDocument(
            doc_id=doc_id,
            text=cleaned_text,
            metadata=metadata,
        )

    def _parse_pdf(self, content_bytes: bytes) -> str:
        reader = pypdf.PdfReader(io.BytesIO(content_bytes))
        pages_text: list[str] = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages_text.append(page_text)
        return "\n\n".join(pages_text)

    def _parse_docx(self, content_bytes: bytes) -> str:
        doc = docx.Document(io.BytesIO(content_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    def _parse_text(self, content_bytes: bytes) -> str:
        try:
            return content_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return content_bytes.decode("latin-1", errors="ignore")
